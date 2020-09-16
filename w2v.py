import sys
import docx
import numpy as np
from nltk.tokenize import word_tokenize, sent_tokenize
from gensim.models import Word2Vec
import logging


def prepareDocument(doc):
    p = doc.paragraphs
    result = list()
    for para in p:
        para = para.text
        sentences = sent_tokenize(para)
        for sent in sentences:
            result.append(word_tokenize(sent))
    return result


def getAverageDocVector(sentences):
    vector = np.zeros(model.vector_size)
    count = 0
    oov = 0
    for sent in sentences:
        words = word_tokenize(sent)
        for word in words:
            try:
                vector += model.wv.get_vector(word)
            except:
                oov += 1
                vector += 0
            count += 1
    vector /= count
    if oov != 0:
        print(f"Out of Vocabulary words: {oov}")
    return vector


def cosine_similarity(A, B):
    return np.dot(A, B)/(np.linalg.norm(A)*np.linalg.norm(B))


filenames = sys.argv[1:]
train_text = list()
for fname in filenames:
    document = docx.Document(fname)
    train_text.extend(prepareDocument(document))

print("Creating Word2Vec model...")
model = Word2Vec(size=500, window=2, iter=20, min_count=1)
model.build_vocab(train_text, progress_per=10)
model.train(train_text, total_examples=model.corpus_count, epochs=50)
print("Training completed...")

print(model.wv.most_similar("persistent"))

print("Obtaining sentences for documents...")
sentences = dict()
for fname in filenames:
    document = docx.Document(fname)
    sentences[fname] = []
    sentences[fname].extend(i.text for i in document.paragraphs)

print("Obtaing Document Vectors...")

vecs = list(map(getAverageDocVector, sentences.values()))

print(f"Similarity = {cosine_similarity(*vecs)}")
