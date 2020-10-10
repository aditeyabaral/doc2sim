import docx
import PyPDF2
from pathlib import Path
import numpy as np
from nltk.tokenize import word_tokenize, sent_tokenize
from gensim.models.doc2vec import Doc2Vec, TaggedDocument

def readDocx(document):
    document = docx.Document(document)
    return ". ".join(para.text.lower() for para in document.paragraphs)


def readPDF(document):
    pdfFileObj = open(document, 'rb')
    pdfReader = PyPDF2.PdfFileReader(pdfFileObj)
    content = ". ".join([pdfReader.getPage(page).extractText().lower()
                         for page in range(pdfReader.numPages)])
    return content


def readOther(document):
    with open(document) as f:
        content = f.read().lower().strip()
    return content


def getText(document):
    p = Path(document)
    extension = str(p.suffix)
    if extension == ".docx":
        return readDocx(document)
    elif extension == ".pdf":
        return readPDF(document)
    else:
        return readOther(document)


def cosine_similarity(A, B):
    return np.dot(A, B)/(np.linalg.norm(A)*np.linalg.norm(B))


def createDoc2VecModel(train_text):
    model = Doc2Vec(vector_size=300, window=2, epochs=20, min_count=1, seed=0)
    model.build_vocab(train_text)
    model.train(train_text, total_examples=model.corpus_count, epochs=50)
    return model

def check_similarity(filenames):
    documents = list(map(getText, filenames))
    train_text = [
        TaggedDocument(
            words=word_tokenize(doc),
            tags=[str(i)]
        ) for i, doc in enumerate(documents)
    ]

    test_text = list(map(word_tokenize, documents))

    #model = createDoc2VecModel(train_text)
    #vecs = list(map(model.infer_vector, test_text))

    total_files = len(filenames)
    similarity_matrix = np.identity(total_files, dtype=float)

    for i in range(total_files):
        for j in range(total_files):
            if i!=j and similarity_matrix[i, j] == 0:
                    model = createDoc2VecModel([train_text[i]] + [train_text[j]])
                    vecs = list(map(model.infer_vector, [test_text[i]] + [test_text[j]]))
                    similarity_matrix[i, j] = cosine_similarity(vecs[0], vecs[1])
                    similarity_matrix[j, i] = cosine_similarity(vecs[0], vecs[1])
                    del model
    return similarity_matrix
